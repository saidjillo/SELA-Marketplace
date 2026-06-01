#!/usr/bin/env python3
"""
PDF → Word (.docx) converter
Usage: python3 convert_pdf.py <input.pdf> <output.docx>
Prints JSON progress updates to stdout.
"""
import sys
import json
import os

def progress(step, total, message):
    print(json.dumps({"step": step, "total": total, "message": message}), flush=True)

def convert(pdf_path, docx_path):
    try:
        from pypdf import PdfReader
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        progress(1, 6, "Opening PDF file…")
        reader = PdfReader(pdf_path)
        total_pages = len(reader.pages)

        if total_pages == 0:
            print(json.dumps({"error": "PDF has no pages"}), flush=True)
            sys.exit(1)

        progress(2, 6, f"Detected {total_pages} page{'s' if total_pages > 1 else ''}…")

        # Build word doc
        doc = Document()

        # Set narrow margins
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        section = doc.sections[0]
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1.2)
        section.right_margin  = Inches(1.2)

        progress(3, 6, "Extracting text from pages…")

        all_lines = []
        for i, page in enumerate(reader.pages):
            raw = page.extract_text()
            if raw:
                lines = raw.split('\n')
                all_lines.extend(lines)
                all_lines.append('')   # blank line between pages
            if (i + 1) % max(1, total_pages // 4) == 0:
                pct = int((i + 1) / total_pages * 100)
                progress(3, 6, f"Extracting text… {pct}% ({i+1}/{total_pages} pages)")

        progress(4, 6, "Formatting document structure…")

        def looks_like_heading(line):
            stripped = line.strip()
            if not stripped:
                return False
            words = stripped.split()
            # Short line, mostly title-cased or ALL CAPS, no sentence ending
            if len(words) <= 10 and (stripped.isupper() or stripped.istitle()):
                return True
            if len(stripped) < 80 and stripped.endswith(':'):
                return True
            return False

        # Write lines to docx with basic structure detection
        for line in all_lines:
            stripped = line.strip()
            if not stripped:
                doc.add_paragraph('')
                continue
            if looks_like_heading(stripped):
                h = doc.add_heading(stripped, level=2)
                h.style.font.color.rgb = RGBColor(0x1a, 0x56, 0x9c)
            else:
                p = doc.add_paragraph(stripped)
                p.style.font.size = Pt(11)

        progress(5, 6, "Saving Word document…")
        doc.save(docx_path)

        # Verify output exists
        size_kb = round(os.path.getsize(docx_path) / 1024, 1)
        progress(6, 6, f"Done! File size: {size_kb} KB")
        print(json.dumps({"success": True, "pages": total_pages, "size_kb": size_kb}), flush=True)

    except Exception as e:
        print(json.dumps({"error": str(e)}), flush=True)
        sys.exit(1)

if __name__ == '__main__':
    if len(sys.argv) != 3:
        print(json.dumps({"error": "Usage: convert_pdf.py <input.pdf> <output.docx>"}))
        sys.exit(1)
    convert(sys.argv[1], sys.argv[2])
